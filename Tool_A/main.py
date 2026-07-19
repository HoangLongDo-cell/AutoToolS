import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import json
import os
import glob
import webbrowser
import pyperclip
from docx import Document
from datetime import datetime
import re

CONFIG_FILE = "config.json"

def clean_corrected_text(text):
    # Chuẩn hóa newline
    text = text.replace('\r\n', '\n')
    
    # Xóa các dạng code block
    lines = text.split('\n')
    cleaned_lines = []
    
    code_block_patterns = [
        r'^```.*$',  # Match bất kỳ dòng nào bắt đầu bằng ```
        r'^`$'       # Match backtick lẻ
    ]
    
    in_code_block = False
    for line in lines:
        if re.match(r'^```', line.strip()):
            continue
        if line.strip() == '`':
            continue
            
        # Xóa các markdown formatting như **, *, __, _, nhưng giữ nguyên text
        # Cẩn thận không xóa các dòng bắt đầu bằng list (vd "1. ")
        
        # Xóa in đậm, in nghiêng (**) và (*) - nhưng không làm mất nội dung
        # Thay vì regex phức tạp, có thể dùng regex thay thế:
        # Xóa ** bao quanh chữ
        l = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        l = re.sub(r'\*(.*?)\*', r'\1', l)
        l = re.sub(r'__(.*?)__', r'\1', l)
        
        # Xóa bullet points (*) ở đầu dòng
        if l.strip().startswith('* ') or l.strip().startswith('- '):
            l = l.strip()[2:]
            
        # Xóa dòng chỉ có *, **, ***, ---, ___
        if re.match(r'^[\*\-\_]{1,}$', l.strip()):
            continue
            
        cleaned_lines.append(l)

    # Rút nhiều dòng trống
    result = []
    empty_count = 0
    for line in cleaned_lines:
        if line.strip() == '':
            empty_count += 1
            if empty_count <= 1:
                result.append(line)
        else:
            empty_count = 0
            result.append(line)
            
    # Xử lý số thứ tự đặc biệt: "1." ở một dòng, dòng tiếp theo là nội dung thì gộp lại
    final_result = []
    i = 0
    while i < len(result):
        curr = result[i]
        if re.match(r'^\d+\.$', curr.strip()) and i + 1 < len(result) and result[i+1].strip() != '':
            final_result.append(f"{curr.strip()} {result[i+1].strip()}")
            i += 2
        else:
            final_result.append(curr)
            i += 1
            
    return '\n'.join(final_result).strip()

class WTPromptTool:
    def __init__(self, root):
        self.root = root
        self.root.title("WT Prompt Tool")
        self.root.geometry("1100x800")
        
        self.config = {
            "input_folder": "D:/WT_INPUT",
            "gpt_url": "",
            "downloads_folder": "D:/HoangLong_Data/Download"
        }
        self.load_config()
        self.ensure_folders()
        
        self.selected_files = []
        
        self.create_widgets()
        self.bind_hotkeys()
        
    def load_config(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    self.config.update(json.load(f))
            except Exception:
                pass
        else:
            self.save_config()
            
    def save_config(self):
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(self.config, f, indent=4, ensure_ascii=False)
            
    def ensure_folders(self):
        os.makedirs(self.config['input_folder'], exist_ok=True)
        try:
            os.makedirs(self.config['downloads_folder'], exist_ok=True)
        except Exception:
            self.config['downloads_folder'] = os.path.join(os.getcwd(), 'outputs')
            os.makedirs(self.config['downloads_folder'], exist_ok=True)
            self.save_config()
            
    def create_widgets(self):
        # Frame chính
        main_paned = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        main_paned.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        left_frame = ttk.Frame(main_paned)
        right_frame = ttk.Frame(main_paned)
        main_paned.add(left_frame, weight=1)
        main_paned.add(right_frame, weight=1)
        
        # === LEFT FRAME ===
        
        # 1. File học viên đã chọn
        lf_file = ttk.LabelFrame(left_frame, text="1. File học viên đã chọn")
        lf_file.pack(fill=tk.X, pady=5)
        
        btn_frame = ttk.Frame(lf_file)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(btn_frame, text="Chọn file", command=self.select_files).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Dùng file mới nhất", command=self.use_latest_file).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Dùng tất cả file", command=self.use_all_files).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Xóa danh sách", command=self.clear_files).pack(side=tk.LEFT, padx=2)
        
        self.file_listbox = tk.Listbox(lf_file, height=4)
        self.file_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 2. Chọn hệ thi
        lf_sys = ttk.LabelFrame(left_frame, text="2. Chọn hệ thi để tạo prompt")
        lf_sys.pack(fill=tk.X, pady=5)
        
        sys_btn_frame = ttk.Frame(lf_sys)
        sys_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(sys_btn_frame, text="APTIS GENERAL (F1)", command=lambda: self.generate_prompt("aptis gen")).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        ttk.Button(sys_btn_frame, text="VSTEP (F2)", command=lambda: self.generate_prompt("vstep")).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        ttk.Button(sys_btn_frame, text="APTIS ADVANCED (F3)", command=lambda: self.generate_prompt("aptis adv")).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        
        # 3. Preview prompt
        lf_preview = ttk.LabelFrame(left_frame, text="3. Preview prompt")
        lf_preview.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.txt_preview = tk.Text(lf_preview, wrap=tk.WORD)
        self.txt_preview.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        prev_btn_frame = ttk.Frame(lf_preview)
        prev_btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Button(prev_btn_frame, text="Copy lại prompt (Ctrl+P)", command=self.copy_prompt).pack(side=tk.LEFT, padx=2)
        ttk.Button(prev_btn_frame, text="Mở GPTs", command=self.open_gpts).pack(side=tk.LEFT, padx=2)
        ttk.Button(prev_btn_frame, text="Lưu .txt", command=self.save_prompt_txt).pack(side=tk.LEFT, padx=2)
        
        # === RIGHT FRAME ===
        
        # 4. Nhắc sửa GPTs - Hay dùng
        lf_quick = ttk.LabelFrame(right_frame, text="4. Nhắc sửa GPTs - Hay dùng")
        lf_quick.pack(fill=tk.X, pady=5)
        
        quick_frame1 = ttk.Frame(lf_quick)
        quick_frame1.pack(fill=tk.X, padx=5, pady=2)
        quick_frame2 = ttk.Frame(lf_quick)
        quick_frame2.pack(fill=tk.X, padx=5, pady=2)
        
        btn_labels = ["Giữ format", "Chỉ sửa lỗi", "Đếm từ lại", "Trả plain text", 
                      "Sai hệ thi", "Sai format chữa lỗi", "Rà lỗi kỹ", "Làm lại Knowledge"]
        for i, lbl in enumerate(btn_labels):
            parent = quick_frame1 if i < 4 else quick_frame2
            ttk.Button(parent, text=lbl, command=lambda t=lbl: self.copy_remind(t)).pack(side=tk.LEFT, fill=tk.X, expand=True, padx=1)
            
        # 5. Nhắc sửa nâng cao
        lf_adv = ttk.LabelFrame(right_frame, text="5. Nhắc sửa nâng cao")
        lf_adv.pack(fill=tk.X, pady=5)
        
        adv_frame = ttk.Frame(lf_adv)
        adv_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self.adv_combo = ttk.Combobox(adv_frame, values=[
            "Xóa keyword", "Không code block", "Dấu chấm ký tên", "Thiếu từ + gợi ý",
            "Không ép thì", "Giải thích TV", "VSTEP T1 không ký", "Hobbies are",
            "Không sửa đề", "Không đổi từ đúng", "Không bỏ sót đoạn", "Không đếm đề",
            "Không gộp file", "ADV góp ý", "Không markdown", "Sửa lỗi nhỏ", "Không overcorrect"
        ], state="readonly")
        self.adv_combo.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=2)
        if self.adv_combo['values']:
            self.adv_combo.current(0)
            
        ttk.Button(adv_frame, text="Copy nhắc sửa", command=lambda: self.copy_remind(self.adv_combo.get())).pack(side=tk.LEFT, padx=2)
        
        # 6. Bài đã chữa từ GPTs
        lf_result = ttk.LabelFrame(right_frame, text="6. Bài đã chữa từ GPTs")
        lf_result.pack(fill=tk.BOTH, expand=True, pady=5)
        
        res_btn_frame1 = ttk.Frame(lf_result)
        res_btn_frame1.pack(fill=tk.X, padx=5, pady=2)
        
        ttk.Button(res_btn_frame1, text="Dán bài đã chữa (Ctrl+D)", command=self.paste_result).pack(side=tk.LEFT, padx=2)
        ttk.Button(res_btn_frame1, text="Xóa nội dung (Ctrl+Del)", command=self.clear_result).pack(side=tk.LEFT, padx=2)
        ttk.Button(res_btn_frame1, text="Định dạng lại (Ctrl+F)", command=self.format_result).pack(side=tk.LEFT, padx=2)
        
        res_btn_frame2 = ttk.Frame(lf_result)
        res_btn_frame2.pack(fill=tk.X, padx=5, pady=2)
        
        ttk.Button(res_btn_frame2, text="Copy bài đã định dạng (Ctrl+Shift+C)", command=self.copy_formatted).pack(side=tk.LEFT, padx=2)
        ttk.Button(res_btn_frame2, text="Xuất Word bài chữa (Ctrl+W)", command=self.export_word).pack(side=tk.LEFT, padx=2)
        ttk.Button(res_btn_frame2, text="Mở thư mục (Ctrl+O)", command=self.open_folder).pack(side=tk.LEFT, padx=2)
        
        self.txt_result = tk.Text(lf_result, wrap=tk.WORD)
        self.txt_result.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 7. Dòng trạng thái
        self.status_var = tk.StringVar()
        self.status_var.set("Sẵn sàng.")
        self.status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def bind_hotkeys(self):
        self.root.bind("<Control-u>", lambda e: self.select_files())
        self.root.bind("<Control-1>", lambda e: self.generate_prompt("aptis gen"))
        self.root.bind("<Control-2>", lambda e: self.generate_prompt("vstep"))
        self.root.bind("<Control-3>", lambda e: self.generate_prompt("aptis adv"))
        
        # Thêm phím tắt F1, F2, F3 để dễ giả lập từ AutoHotkey (tránh lỗi Unikey)
        self.root.bind("<F1>", lambda e: self.generate_prompt("aptis gen"))
        self.root.bind("<F2>", lambda e: self.generate_prompt("vstep"))
        self.root.bind("<F3>", lambda e: self.generate_prompt("aptis adv"))
        
        self.root.bind("<Control-p>", lambda e: self.copy_prompt())
        self.root.bind("<Control-d>", lambda e: self.paste_result())
        self.root.bind("<Control-f>", lambda e: self.format_result())
        self.root.bind("<Control-Shift-C>", lambda e: self.copy_formatted())
        self.root.bind("<Control-w>", lambda e: self.export_word())
        self.root.bind("<Control-o>", lambda e: self.open_folder())
        self.root.bind("<Control-Delete>", lambda e: self.clear_result())

    def update_status(self, msg):
        self.status_var.set(msg)
        
    def read_file_content(self, filepath):
        ext = os.path.splitext(filepath)[1].lower()
        if ext == '.txt':
            for enc in ['utf-8-sig', 'utf-8', 'cp1258', 'cp1252']:
                try:
                    with open(filepath, 'r', encoding=enc) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return ""
        elif ext in ['.docx', '.doc', '.rtf']:
            if ext == '.docx':
                try:
                    doc = Document(filepath)
                    text = []
                    for para in doc.paragraphs:
                        text.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text.append(cell.text)
                    return '\n'.join(text)
                except Exception:
                    pass
                    
            try:
                import win32com.client
                import pythoncom
                pythoncom.CoInitialize()
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                doc = None
                text = ""
                try:
                    abs_path = os.path.abspath(filepath)
                    doc = word.Documents.Open(abs_path, ReadOnly=True, Visible=False)
                    text = doc.Content.Text
                    text = text.replace('\r', '\n')
                finally:
                    if doc is not None:
                        doc.Close(False)
                    word.Quit()
                return text
            except Exception as e:
                return f"Lỗi đọc file Word: {e}"
        return ""

    def select_files(self):
        files = filedialog.askopenfilenames(
            initialdir=self.config['downloads_folder'],
            title="Chọn file học viên",
            filetypes=(("Text/Word files", "*.txt *.docx *.doc"), ("All files", "*.*"))
        )
        if files:
            self.selected_files = list(files)
            self.update_file_listbox()
            self.update_status(f"Đã chọn {len(self.selected_files)} file.")
            
        self.root.focus_force()

    def use_latest_file(self):
        folder = self.config['input_folder']
        files = glob.glob(os.path.join(folder, "*.txt")) + glob.glob(os.path.join(folder, "*.docx")) + glob.glob(os.path.join(folder, "*.doc"))
        if not files:
            self.update_status("Không tìm thấy file nào trong thư mục input.")
            return
        latest_file = max(files, key=os.path.getctime)
        self.selected_files = [latest_file]
        self.update_file_listbox()
        self.update_status("Đã chọn file mới nhất.")

    def use_all_files(self):
        folder = self.config['input_folder']
        files = glob.glob(os.path.join(folder, "*.txt")) + glob.glob(os.path.join(folder, "*.docx")) + glob.glob(os.path.join(folder, "*.doc"))
        files.sort()
        if not files:
            self.update_status("Không tìm thấy file nào trong thư mục input.")
            return
        self.selected_files = files
        self.update_file_listbox()
        self.update_status(f"Đã chọn tất cả {len(self.selected_files)} file.")

    def clear_files(self):
        self.selected_files = []
        self.update_file_listbox()
        self.update_status("Đã xóa danh sách file.")

    def update_file_listbox(self):
        self.file_listbox.delete(0, tk.END)
        for f in self.selected_files:
            self.file_listbox.insert(tk.END, os.path.basename(f))

    def generate_prompt(self, keyword):
        if not self.selected_files:
            self.update_status("Chưa chọn file nào để tạo prompt!")
            return
            
        prompt = f"[{keyword}]\n\n"
        for f in self.selected_files:
            content = self.read_file_content(f)
            prompt += f"{content.strip()}\n\n\n"
            
        self.txt_preview.delete("1.0", tk.END)
        self.txt_preview.insert(tk.END, prompt.strip())
        self.update_status(f"Đã tạo prompt hệ: {keyword.upper()}")

    def copy_prompt(self):
        content = self.txt_preview.get("1.0", tk.END).strip()
        pyperclip.copy(content)
        self.update_status("Đã copy prompt vào clipboard.")

    def open_gpts(self):
        url = self.config.get('gpt_url', '')
        if url:
            webbrowser.open(url)
            self.update_status("Đã mở GPTs.")
        else:
            self.update_status("Chưa cấu hình link GPTs.")

    def save_prompt_txt(self):
        content = self.txt_preview.get("1.0", tk.END).strip()
        path = os.path.join(self.config['downloads_folder'], "prompt.txt")
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        self.update_status(f"Đã lưu prompt ra file {path}")

    def copy_remind(self, text):
        pyperclip.copy(text)
        self.update_status(f"Đã copy nhắc sửa: {text}")

    def paste_result(self):
        try:
            content = pyperclip.paste()
            self.txt_result.delete("1.0", tk.END)
            self.txt_result.insert(tk.END, content)
            self.update_status("Đã dán bài chữa từ clipboard.")
        except Exception:
            self.update_status("Lỗi khi dán từ clipboard.")

    def clear_result(self):
        self.txt_result.delete("1.0", tk.END)
        self.update_status("Đã xóa nội dung bài chữa.")

    def format_result(self):
        content = self.txt_result.get("1.0", tk.END).strip()
        cleaned = clean_corrected_text(content)
        self.txt_result.delete("1.0", tk.END)
        self.txt_result.insert(tk.END, cleaned)
        self.update_status("Đã định dạng lại bài chữa.")
        
    def copy_formatted(self):
        self.format_result()
        cleaned = self.txt_result.get("1.0", tk.END).strip()
        pyperclip.copy(cleaned)
        self.update_status("Đã copy bài đã định dạng vào clipboard.")

    def export_word(self):
        self.format_result()
        content = self.txt_result.get("1.0", tk.END).strip()
        if not content:
            self.update_status("Không có nội dung để xuất Word!")
            return
            
        doc = Document()
        # Set font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        
        for line in content.split('\n'):
            doc.add_paragraph(line)
            
        # Tìm tên file phù hợp
        base_name = "Bài chữa"
        ext = ".docx"
        out_folder = self.config['downloads_folder']
        out_path = os.path.join(out_folder, base_name + ext)
        
        counter = 1
        while os.path.exists(out_path):
            out_path = os.path.join(out_folder, f"{base_name} ({counter}){ext}")
            counter += 1
            
        try:
            doc.save(out_path)
            self.update_status(f"Đã lưu file: {out_path}")
        except Exception as e:
            self.update_status(f"Lỗi khi lưu file Word: {e}")

    def open_folder(self):
        folder = self.config['downloads_folder']
        if os.path.exists(folder):
            os.startfile(folder)
            self.update_status("Đã mở thư mục bài chữa.")
        else:
            self.update_status("Thư mục bài chữa không tồn tại.")

if __name__ == "__main__":
    root = tk.Tk()
    app = WTPromptTool(root)
    root.mainloop()
